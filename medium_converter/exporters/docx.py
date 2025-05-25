"""DOCX exporter for Medium articles."""

from typing import BinaryIO, TextIO
import io

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from ..core.models import Article, ContentBlock, ContentType, Section
from .base import BaseExporter


class DocxExporter(BaseExporter):
    """Export Medium articles to DOCX format."""

    def __init__(self):
        """Initialize the DOCX exporter."""
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install medium-converter[word]"
            )
        super().__init__()

    def export(
        self, article: Article, output: str | TextIO | BinaryIO | None = None
    ) -> bytes:
        """Export an article to DOCX.

        Args:
            article: The article to export
            output: Optional output file path or file-like object

        Returns:
            The exported content as bytes
        """
        # Create document
        doc = docx.Document()

        # Set document metadata
        doc.core_properties.title = article.title
        doc.core_properties.author = article.author

        # Add title
        title = doc.add_heading(article.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add author and date
        byline = doc.add_paragraph(f"By {article.author} | {article.date}")
        byline.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add tags if available
        if article.tags:
            tags = ", ".join([f"#{tag.replace(' ', '')}" for tag in article.tags])
            tag_paragraph = doc.add_paragraph(tags)
            tag_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add reading time if available
        if article.estimated_reading_time:
            reading_time = doc.add_paragraph(
                f"{article.estimated_reading_time} min read"
            )
            reading_time.italic = True
            reading_time.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a divider
        doc.add_paragraph("").add_run().add_break()

        # Process content
        for item in article.content:
            if isinstance(item, Section):
                if item.title:
                    doc.add_heading(item.title, 1)

                for block in item.blocks:
                    self._format_block(doc, block)
            elif isinstance(item, ContentBlock):
                self._format_block(doc, block)

        # Save document
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        if output:
            if isinstance(output, str):
                doc.save(output)
            else:
                # If it's a file-like object
                if hasattr(output, "write") and callable(output.write):
                    output.write(doc_bytes.read())
                    doc_bytes.seek(0)  # Reset for the return value

        return doc_bytes.read()

    def _format_block(self, doc, block: ContentBlock) -> None:
        """Format a content block in DOCX format.

        Args:
            doc: The docx Document object
            block: The content block to format
        """
        if block.type == ContentType.TEXT:
            p = doc.add_paragraph(block.content)

        elif block.type == ContentType.HEADING:
            level = block.metadata.get("level", 2)
            # Ensure level is valid (1-9)
            level = max(1, min(9, level))
            doc.add_heading(block.content, level)

        elif block.type == ContentType.IMAGE:
            try:
                p = doc.add_paragraph()
                r = p.add_run()
                # Note: This only works for local images
                # For remote images, we'd need to download them first
                img_path = block.content
                if img_path.startswith(("http://", "https://")):
                    # Skip remote images for simplicity
                    # In a real implementation, we'd download these
                    r.add_text(f"[Image: {block.metadata.get('alt', 'Image')}]")
                else:
                    r.add_picture(img_path, width=Inches(6))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add caption if available
                if "alt" in block.metadata:
                    caption = doc.add_paragraph(block.metadata["alt"])
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = "Caption"
            except Exception:
                # If image insertion fails, add a placeholder
                p = doc.add_paragraph(f"[Image: {block.metadata.get('alt', 'Image')}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif block.type == ContentType.CODE:
            p = doc.add_paragraph()
            code_text = f"{block.content}"
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        elif block.type == ContentType.QUOTE:
            p = doc.add_paragraph()
            p.style = "Quote"
            p.add_run(block.content)

        elif block.type == ContentType.LIST:
            list_type = block.metadata.get("list_type", "unordered")
            items = block.content.split("\n")

            for item in items:
                if not item.strip():
                    continue

                p = doc.add_paragraph(
                    style="List Bullet" if list_type == "unordered" else "List Number"
                )
                p.add_run(item)
        else:
            # Default case
            doc.add_paragraph(block.content)
